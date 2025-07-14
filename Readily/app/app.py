from flask import Flask, render_template, request, send_file
import requests
from bs4 import BeautifulSoup
from docx import Document
from fpdf import FPDF
import io

app = Flask(__name__)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        url = request.form['url']
        filetype = request.form.get('filetype', 'docx')
        try:
            response = requests.get(url)
            soup = BeautifulSoup(response.text, 'html.parser')
            # Try to extract main content
            paragraphs = soup.find_all('p')
            text = '\n\n'.join([p.get_text() for p in paragraphs if p.get_text(strip=True)])
            if not text.strip():
                text = 'No readable content found.'
            if filetype == 'pdf':
                # Replace curly quotes and unsupported characters with ASCII equivalents
                safe_text = text.replace('’', "'").replace('‘', "'").replace('“', '"').replace('”', '"')
                safe_text = safe_text.encode('latin-1', 'replace').decode('latin-1')
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font('helvetica', size=12)
                pdf.multi_cell(0, 10, f'Content from {url}\n\n{safe_text}')
                file_stream = io.BytesIO()
                pdf.output(file_stream)
                file_stream.seek(0)
                return send_file(file_stream, as_attachment=True, download_name='web_content.pdf', mimetype='application/pdf')
            else:
                doc = Document()
                doc.add_heading(f'Content from {url}', 0)
                doc.add_paragraph(text)
                file_stream = io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                return send_file(file_stream, as_attachment=True, download_name='web_content.docx')
        except Exception as e:
            return f"Error: {e}", 500
    return render_template('index.html')

if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True)
